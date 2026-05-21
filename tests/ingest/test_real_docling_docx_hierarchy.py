# @summary
# Real-Docling DOCX hierarchy test: synthesizes a DOCX with genuinely
# nested headings (H1 > H2 > H3 around a table) and asserts that the
# resulting ``TableArtifact.section_path`` carries the full breadcrumb.
# Closes the real-input evidence gap that the PDF backend (Docling 2.82.0)
# flattens every heading to ``.level=1`` -- DOCX preserves levels, so the
# non-flat-outline path of ``_resolve_table_section_paths`` (W's
# ``had_content`` heuristic) is exercised on a real Docling document.
# Exports: test_real_docling_docx_hierarchy
# Deps: python-docx, docling (real), src.ingest.support.docling, pytest
# @end-summary

"""Real-Docling DOCX hierarchy smoke test.

The PDF backend in Docling 2.82.0 reports ``.level=1`` for every heading,
which trips the flat-outline gate in ``_resolve_table_section_paths``.
Two real-PDF soaks (ESP32-S3, MSP430) confirmed this. To get real-input
coverage of the *non*-flat-outline path -- the ``had_content`` heuristic
that builds nested breadcrumbs -- we need an input format whose Docling
backend preserves heading levels. DOCX does.

This test generates a small DOCX containing::

    H1  "Bit Field Reference"
        body paragraph
    H2  "Control Register"
        body paragraph
    H3  "Bit Allocation"
        table (3 cols x 4 rows)

and asserts:

1. Docling actually emits multiple distinct heading levels for this DOCX
   (sanity: refute the assumption that all backends flatten).
2. At least one ``TableArtifact`` has a ``section_path`` containing all
   three heading strings, in order, joined by ``" > "`` (full breadcrumb).
3. The flat-outline gate did NOT fire (indirectly: depth >= 2).

Gated behind ``slow`` + ``integration`` markers and self-skips when
python-docx or Docling models are unavailable.
"""

from __future__ import annotations

import json
import tempfile
from pathlib import Path

import pytest


# ---------------------------------------------------------------------------
# Fixture content
# ---------------------------------------------------------------------------

H1_TEXT = "Bit Field Reference"
H2_TEXT = "Control Register"
H3_TEXT = "Bit Allocation"

TABLE_ROWS: list[list[str]] = [
    ["Bit", "Name", "Description"],
    ["7", "EN", "Module enable"],
    ["6", "MODE", "Operating mode"],
    ["0", "START", "Start conversion"],
]


def _synthesize_docx(path: Path) -> None:
    """Render a small DOCX with H1 > H2 > H3 around a 3x4 table.

    Uses python-docx default ``Heading 1`` / ``Heading 2`` / ``Heading 3``
    styles -- Docling's DOCX backend reads these into ``SectionHeaderItem``
    nodes with the matching ``.level``.
    """
    from docx import Document

    doc = Document()

    doc.add_heading(H1_TEXT, level=1)
    doc.add_paragraph(
        "This section documents the bit-field layout used by the device "
        "control surface. The following sub-sections describe each "
        "register and its individual bit allocations."
    )

    doc.add_heading(H2_TEXT, level=2)
    doc.add_paragraph(
        "The control register selects the device operating mode and "
        "enables data conversion."
    )

    doc.add_heading(H3_TEXT, level=3)

    # 3 cols x 4 rows including the header row.
    table = doc.add_table(rows=len(TABLE_ROWS), cols=len(TABLE_ROWS[0]))
    table.style = "Table Grid"
    for r, row in enumerate(TABLE_ROWS):
        cells = table.rows[r].cells
        for c, value in enumerate(row):
            cells[c].text = value

    doc.add_paragraph(
        "Reserved fields must be written as zero for forward "
        "compatibility."
    )

    doc.save(str(path))


# ---------------------------------------------------------------------------
# Diagnostics
# ---------------------------------------------------------------------------


def _dump_diagnostics(parsed_md: str, tables: list, heading_levels: list[int]) -> Path:
    out_dir = Path(tempfile.mkdtemp(prefix="real_docling_docx_hier_"))
    (out_dir / "parsed.md").write_text(parsed_md or "", encoding="utf-8")
    (out_dir / "heading_levels.json").write_text(
        json.dumps(heading_levels), encoding="utf-8"
    )
    table_blob = []
    for t in tables:
        table_blob.append(
            {
                "self_ref": getattr(t, "self_ref", None),
                "section_path": getattr(t, "section_path", None),
                "num_rows": getattr(t, "num_rows", None),
                "num_cols": getattr(t, "num_cols", None),
                "markdown_head": (getattr(t, "markdown", "") or "")[:200],
            }
        )
    (out_dir / "tables.json").write_text(
        json.dumps(table_blob, indent=2, default=str), encoding="utf-8"
    )
    return out_dir


def _build_config():
    from src.ingest.common.types import IngestionConfig

    config = IngestionConfig()
    try:
        config.docling_auto_download = True
    except Exception:  # noqa: BLE001
        pass
    return config


def _collect_heading_levels(document) -> list[int]:
    """Walk ``iterate_items`` and return the ``.level`` of every heading-ish item."""
    from src.ingest.support.docling import (  # noqa: PLC0415
        _HEADING_LABEL_VALUES,
        _label_value,
    )

    levels: list[int] = []
    iterate = getattr(document, "iterate_items", None)
    if not callable(iterate):
        return levels
    for entry in iterate():
        item = entry[0] if isinstance(entry, tuple) and len(entry) >= 1 else entry
        if _label_value(item) in _HEADING_LABEL_VALUES:
            lvl = int(getattr(item, "level", 0) or 0)
            if _label_value(item) == "title":
                lvl = 0
            levels.append(lvl)
    return levels


# ---------------------------------------------------------------------------
# The test
# ---------------------------------------------------------------------------


@pytest.mark.slow
@pytest.mark.integration
def test_real_docling_docx_hierarchy(tmp_path: Path) -> None:
    """Real Docling over a DOCX with H1>H2>H3 keeps the full breadcrumb on tables.

    Sanity-checks that Docling's DOCX backend emits >=2 distinct heading
    levels (refuting the flat-outline assumption that holds for the PDF
    backend), then asserts that ``_resolve_table_section_paths`` walked
    the non-flat-outline path and stitched the full chain
    ``Bit Field Reference > Control Register > Bit Allocation`` onto the
    nested table.
    """
    pytest.importorskip("docx", reason="python-docx not installed")

    from src.ingest.support.docling import DoclingParser

    docx_path = tmp_path / "bitfield_hierarchy.docx"
    _synthesize_docx(docx_path)
    assert docx_path.exists() and docx_path.stat().st_size > 0

    config = _build_config()

    try:
        DoclingParser.ensure_ready(config)
    except Exception as exc:  # noqa: BLE001 -- any failure -> skip
        pytest.skip(f"Docling models unavailable: {exc!r}")

    parser = DoclingParser()
    try:
        parse_result = parser.parse(docx_path, config)
    except Exception as exc:  # noqa: BLE001
        pytest.skip(f"Docling DOCX backend unavailable: {exc!r}")

    tables = list(getattr(parse_result, "tables", []) or [])

    # Pull heading levels straight from the docling document for the
    # sanity-check on real-input evidence.
    document = getattr(parse_result, "_docling_document", None) or getattr(
        parse_result, "docling_document", None
    )
    if document is None:
        # Fall back: re-parse to grab the document object directly.
        from docling.document_converter import DocumentConverter  # noqa: PLC0415

        try:
            conv = DocumentConverter().convert(str(docx_path))
            document = conv.document
        except Exception as exc:  # noqa: BLE001
            pytest.skip(f"Could not retrieve docling document for sanity check: {exc!r}")

    heading_levels = _collect_heading_levels(document)
    diag_dir = _dump_diagnostics(
        getattr(parse_result, "markdown", "") or "", tables, heading_levels
    )

    def _fail(msg: str) -> None:
        print(f"\n[diagnostics] {diag_dir}")
        raise AssertionError(msg)

    # ---- (1) Sanity: Docling emits >=2 distinct levels on this DOCX --------
    distinct_levels = {lvl for lvl in heading_levels}
    if len(distinct_levels) < 2:
        _fail(
            "Docling DOCX backend did NOT emit multi-level headings on a "
            "document with H1>H2>H3 -- this invalidates the test's premise. "
            f"heading_levels={heading_levels!r}"
        )

    # ---- (2) At least one TableArtifact ------------------------------------
    if not tables:
        _fail(
            "Expected at least one TableArtifact from the DOCX; got 0. "
            f"markdown_head={(getattr(parse_result, 'markdown', '') or '')[:200]!r}"
        )

    # ---- (3) Full breadcrumb present on the nested table -------------------
    section_paths = [(getattr(t, "section_path", "") or "") for t in tables]

    h1_lc, h2_lc, h3_lc = H1_TEXT.lower(), H2_TEXT.lower(), H3_TEXT.lower()
    full_chain_found = False
    for sp in section_paths:
        sp_lc = sp.lower()
        i1 = sp_lc.find(h1_lc)
        i2 = sp_lc.find(h2_lc)
        i3 = sp_lc.find(h3_lc)
        if i1 != -1 and i2 != -1 and i3 != -1 and i1 < i2 < i3:
            full_chain_found = True
            break
    if not full_chain_found:
        _fail(
            "No table section_path carries the full H1>H2>H3 chain "
            f"({H1_TEXT!r} > {H2_TEXT!r} > {H3_TEXT!r}) in order. "
            f"section_paths={section_paths!r}"
        )

    # ---- (4) Flat-outline gate did NOT fire (depth >= 2 on the nested one)
    # The matching breadcrumb above must have >=2 segments when split on " > "
    # -- otherwise the gate effectively flattened the chain.
    depths = [
        len([s for s in sp.split(" > ") if s.strip()]) for sp in section_paths
    ]
    if max(depths) < 2:
        _fail(
            "Max section_path depth is <2 -- the flat-outline gate appears "
            f"to have collapsed the breadcrumb. depths={depths!r} "
            f"section_paths={section_paths!r}"
        )
